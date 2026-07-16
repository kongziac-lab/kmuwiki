import io
import sys
import types
import unittest
from unittest.mock import patch

from kmu_ingest.parsers import extract_text


class TestMultimodalParsers(unittest.TestCase):
    def test_legacy_doc_uses_hardened_word_conversion(self):
        calls = []

        class _Document:
            def SaveAs2(self, path, FileFormat):
                calls.append(("save", FileFormat))
                from docx import Document

                converted = Document()
                converted.add_paragraph("레거시 문서 본문")
                converted.save(path)

            def Close(self, save):
                calls.append(("close", save))

        class _Documents:
            def Open(self, path, **kwargs):
                calls.append(("open", kwargs))
                return _Document()

        class _Word:
            Documents = _Documents()

            def Quit(self):
                calls.append(("quit",))

        word = _Word()
        client = types.SimpleNamespace(DispatchEx=lambda name: word)
        pythoncom = types.SimpleNamespace(CoInitialize=lambda: None, CoUninitialize=lambda: None)
        modules = {
            "pythoncom": pythoncom,
            "win32com": types.SimpleNamespace(client=client),
            "win32com.client": client,
        }

        with patch("kmu_ingest.parsers._word_com_supported", return_value=True), patch.dict(sys.modules, modules):
            parsed = extract_text("legacy.doc", b"binary-doc")

        self.assertIn("레거시 문서 본문", parsed.text)
        self.assertEqual(word.AutomationSecurity, 3)
        self.assertEqual(calls[0][0], "open")
        self.assertTrue(calls[0][1]["ReadOnly"])
        self.assertFalse(calls[0][1]["AddToRecentFiles"])
        self.assertIn(("save", 16), calls)

    def test_pdf_emits_rendered_page_with_page_number(self):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF unavailable")
        document = fitz.open()
        page = document.new_page(width=300, height=200)
        page.insert_text((30, 50), "KMU visual search page")
        payload = document.tobytes()
        document.close()

        result = extract_text("report.pdf", payload)

        pages = [asset for asset in result.assets if asset.asset_type == "page"]
        self.assertEqual(len(pages), 1)
        self.assertEqual(pages[0].page_no, 1)
        self.assertEqual(pages[0].media_type, "image/png")
        self.assertGreater(len(pages[0].image_bytes or b""), 100)

    def test_docx_table_becomes_markdown_search_asset(self):
        import docx

        document = docx.Document()
        document.add_paragraph("예산 보고")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "항목"
        table.cell(0, 1).text = "금액"
        table.cell(1, 0).text = "홍보비"
        table.cell(1, 1).text = "120만원"
        output = io.BytesIO()
        document.save(output)

        result = extract_text("budget.docx", output.getvalue())

        tables = [asset for asset in result.assets if asset.asset_type == "table"]
        self.assertEqual(len(tables), 1)
        self.assertIn("홍보비", tables[0].structured_content)
        self.assertIn("120만원", result.text)

    def test_xlsx_worksheet_becomes_structured_search_asset(self):
        import openpyxl

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "예산"
        sheet.append(["항목", "금액"])
        sheet.append(["홍보비", 1200000])
        output = io.BytesIO()
        workbook.save(output)

        result = extract_text("budget.xlsx", output.getvalue())

        sheets = [asset for asset in result.assets if asset.asset_type == "worksheet"]
        self.assertEqual(len(sheets), 1)
        self.assertEqual(sheets[0].caption, "예산")
        self.assertIn("1200000", sheets[0].structured_content)


if __name__ == "__main__":
    unittest.main()
