"""문서 → 텍스트 추출. 무거운 의존성은 lazy import.

반환 ParseResult.needs_ocr=True 이면 스캔본으로 보고 OCR 단계로 넘긴다.
잠긴 파일은 이 단계에 오지 않는다(lockdetect 에서 걸러짐, 불변식 2).

지원 포맷(실제 전자결재 ZIP 기준): pdf, hwp/hwpx, doc/docx, xls/xlsx, html/htm, mht, txt/csv/md/json/xml.
"""

from __future__ import annotations

import html as _htmlmod
import io
import mimetypes
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from .models import ParsedAsset


@dataclass
class ParseResult:
    text: str
    needs_ocr: bool = False
    mime_type: str | None = None
    assets: list[ParsedAsset] | None = None

    def __post_init__(self) -> None:
        if self.assets is None:
            self.assets = []


_PLAIN_EXT = (".txt", ".csv", ".md", ".json")
_HTML_EXT = (".html", ".htm", ".xml")
_IMG_EXT = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")


def extract_text(filename: str, data: bytes, *, max_visual_pages: int = 100) -> ParseResult:
    name = filename.lower()

    if name.endswith(_PLAIN_EXT):
        return ParseResult(_decode(data), mime_type="text/plain")
    if name.endswith(_HTML_EXT):
        return ParseResult(_strip_html(_decode(data)), mime_type="text/html")
    if name.endswith(".pdf"):
        return _pdf(data, max_visual_pages=max_visual_pages)
    if name.endswith(".docx"):
        return _docx(data)
    if name.endswith(".doc"):
        return _doc(data)
    if name.endswith(".xlsx"):
        return _xlsx(data)
    if name.endswith(".xls"):
        return _xls(data)
    if name.endswith(".mht") or name.endswith(".mhtml"):
        return _mht(data)
    if name.endswith((".hwp", ".hwpx")):
        return _hwp(name, data)
    if name.endswith(_IMG_EXT):
        media_type = _media_type(filename)
        width, height = _image_size(data)
        return ParseResult(
            "",
            needs_ocr=True,
            mime_type=media_type,
            assets=[ParsedAsset(
                asset_type="image",
                page_no=1,
                image_bytes=data,
                media_type=media_type,
                width=width,
                height=height,
                needs_ocr=True,
                extraction_model="source-image",
                extraction_version="v2",
            )],
        )
    return ParseResult("", needs_ocr=False, mime_type=None)


# ── 텍스트/HTML ────────────────────────────────────────────────
def _decode(data: bytes) -> str:
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _strip_html(s: str) -> str:
    s = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", s)
    s = re.sub(r"<[^>]+>", " ", s)
    s = _htmlmod.unescape(s)
    s = re.sub(r"[ \t]+", " ", s)
    return re.sub(r"\n\s*\n+", "\n", s).strip()


# ── PDF ────────────────────────────────────────────────────────
def _pdf(data: bytes, *, max_visual_pages: int = 100) -> ParseResult:
    rendered = _render_pdf_pages(data, max_pages=max_visual_pages)
    try:
        import pdfplumber  # lazy
    except ImportError:
        return ParseResult(
            "", needs_ocr=True, mime_type="application/pdf",
            assets=_rendered_page_assets(rendered, needs_ocr=True),
        )
    try:
        parts: list[str] = []
        assets: list[ParsedAsset] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_no, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""
                parts.append(page_text)
                page_render = rendered.get(page_no)
                assets.append(ParsedAsset(
                    asset_type="page",
                    page_no=page_no,
                    text=page_text,
                    image_bytes=page_render[0] if page_render else None,
                    media_type="image/png" if page_render else None,
                    width=page_render[1] if page_render else None,
                    height=page_render[2] if page_render else None,
                    needs_ocr=not bool(page_text.strip()),
                    extraction_model="pdfplumber+pymupdf",
                    extraction_version="v2",
                ))
                assets.extend(_pdf_table_assets(page, page_no, page_render))
        text = "\n".join(parts).strip()
    except Exception as exc:
        # pdfminer 가 깨지는 PDF(예: "Invalid octal")는 pypdf 로 한 번 더 시도한다.
        # 폴백까지 실패하면 원 예외를 그대로 올려 기존 failed 기록을 유지한다.
        text = _pdf_via_pypdf(data, original=exc)
        assets = _rendered_page_assets(rendered, needs_ocr=(len(text) < 20))
    # 텍스트가 거의 없으면 스캔 PDF로 보고 OCR 대상
    return ParseResult(
        text,
        needs_ocr=(len(text) < 20),
        mime_type="application/pdf",
        assets=assets,
    )


def _render_pdf_pages(data: bytes, *, max_pages: int) -> dict[int, tuple[bytes, int, int]]:
    """Render bounded PDF pages locally; failure never discards extracted text."""
    if max_pages < 1:
        return {}
    try:
        import fitz  # PyMuPDF, lazy

        out: dict[int, tuple[bytes, int, int]] = {}
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page_no, page in enumerate(doc, start=1):
                if page_no > max_pages:
                    break
                pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                out[page_no] = (pix.tobytes("png"), pix.width, pix.height)
        return out
    except Exception:
        return {}


def _rendered_page_assets(
    rendered: dict[int, tuple[bytes, int, int]], *, needs_ocr: bool,
) -> list[ParsedAsset]:
    return [ParsedAsset(
        asset_type="page",
        page_no=page_no,
        image_bytes=payload,
        media_type="image/png",
        width=width,
        height=height,
        needs_ocr=needs_ocr,
        extraction_model="pymupdf",
        extraction_version="v2",
    ) for page_no, (payload, width, height) in sorted(rendered.items())]


def _pdf_table_assets(page, page_no: int, rendered) -> list[ParsedAsset]:
    """Extract table structure and, when possible, the matching visual crop."""
    try:
        tables = page.find_tables() or []
    except Exception:
        return []
    assets: list[ParsedAsset] = []
    for table_no, table in enumerate(tables, start=1):
        try:
            rows = table.extract() or []
            markdown = _table_to_markdown(rows, title=f"페이지 {page_no} 표 {table_no}")
            bbox = tuple(float(value) for value in table.bbox)
            image_bytes = None
            width = height = None
            if rendered:
                image_bytes, width, height = _crop_pdf_render(
                    rendered, bbox, float(page.width), float(page.height))
            assets.append(ParsedAsset(
                asset_type="table",
                page_no=page_no,
                bbox=bbox,
                structured_content=markdown,
                image_bytes=image_bytes,
                media_type="image/png" if image_bytes else None,
                width=width,
                height=height,
                extraction_model="pdfplumber-table",
                extraction_version="v2",
            ))
        except Exception:
            continue
    return assets


def _crop_pdf_render(
    rendered: tuple[bytes, int, int],
    bbox: tuple[float, float, float, float],
    page_width: float,
    page_height: float,
) -> tuple[bytes | None, int | None, int | None]:
    try:
        from PIL import Image

        payload, width, height = rendered
        image = Image.open(io.BytesIO(payload))
        x0, top, x1, bottom = bbox
        box = (
            max(0, round(x0 / page_width * width)),
            max(0, round(top / page_height * height)),
            min(width, round(x1 / page_width * width)),
            min(height, round(bottom / page_height * height)),
        )
        crop = image.crop(box)
        if crop.width < 2 or crop.height < 2:
            return None, None, None
        buf = io.BytesIO()
        crop.save(buf, format="PNG", optimize=True)
        return buf.getvalue(), crop.width, crop.height
    except Exception:
        return None, None, None


def _pdf_via_pypdf(data: bytes, *, original: Exception) -> str:
    try:
        from pypdf import PdfReader  # lazy
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception:
        raise original


# ── Office ─────────────────────────────────────────────────────
def _docx(data: bytes) -> ParseResult:
    import docx  # python-docx, lazy

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    assets: list[ParsedAsset] = []
    for table_no, table in enumerate(d.tables, start=1):
        rows = [[c.text for c in row.cells] for row in table.rows]
        markdown = _table_to_markdown(rows, title=f"표 {table_no}")
        parts.append(markdown)
        assets.append(ParsedAsset(
            asset_type="table",
            structured_content=markdown,
            extraction_model="python-docx",
            extraction_version="v2",
        ))
    assets.extend(_zip_media_assets(data, prefix="word/media/"))
    return ParseResult(
        "\n".join(parts).strip(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        assets=assets,
    )


def _doc(data: bytes) -> ParseResult:
    if not _word_com_supported():
        return ParseResult("", mime_type="application/msword")

    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return ParseResult("", mime_type="application/msword")

    word = None
    document = None
    pythoncom.CoInitialize()
    try:
        with tempfile.TemporaryDirectory(prefix="kmuwiki-doc-") as temp_dir:
            source = Path(temp_dir) / "source.doc"
            converted = Path(temp_dir) / "converted.docx"
            source.write_bytes(data)

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            word.AutomationSecurity = 3
            document = word.Documents.Open(
                str(source),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )
            document.SaveAs2(str(converted), FileFormat=16)
            document.Close(False)
            document = None
            return _docx(converted.read_bytes())
    except Exception:
        return ParseResult("", mime_type="application/msword")
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _word_com_supported() -> bool:
    return os.name == "nt"


def _xlsx(data: bytes) -> ParseResult:
    import openpyxl  # lazy

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=False, data_only=True)
    parts: list[str] = []
    assets: list[ParsedAsset] = []
    for ws in wb.worksheets:
        rows = [["" if v is None else str(v) for v in row]
                for row in ws.iter_rows(values_only=True)]
        markdown = _table_to_markdown(rows, title=ws.title)
        parts.append(markdown)
        assets.append(ParsedAsset(
            asset_type="worksheet",
            structured_content=markdown,
            caption=ws.title,
            extraction_model="openpyxl",
            extraction_version="v2",
        ))
        for chart_no, chart in enumerate(getattr(ws, "_charts", []), start=1):
            caption = _chart_caption(ws.title, chart, chart_no)
            parts.append(caption)
            assets.append(ParsedAsset(
                asset_type="chart",
                structured_content=caption,
                caption=caption,
                extraction_model="openpyxl-chart",
                extraction_version="v2",
            ))
    assets.extend(_zip_media_assets(data, prefix="xl/media/"))
    return ParseResult(
        "\n\n".join(p for p in parts if p.strip()).strip(),
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        assets=assets,
    )


def _xls(data: bytes) -> ParseResult:
    """구형 .xls (BIFF). xlrd 필요."""
    import xlrd  # lazy

    wb = xlrd.open_workbook(file_contents=data)
    parts: list[str] = []
    assets: list[ParsedAsset] = []
    for sh in wb.sheets():
        rows = [[_xls_cell(c) for c in sh.row(r)] for r in range(sh.nrows)]
        markdown = _table_to_markdown(rows, title=sh.name)
        parts.append(markdown)
        assets.append(ParsedAsset(
            asset_type="worksheet",
            structured_content=markdown,
            caption=sh.name,
            extraction_model="xlrd",
            extraction_version="v2",
        ))
    return ParseResult(
        "\n\n".join(parts).strip(),
        mime_type="application/vnd.ms-excel",
        assets=assets,
    )


def _xls_cell(cell) -> str:
    v = cell.value
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return "" if v is None else str(v)


# ── MHT (MHTML 웹 아카이브) ─────────────────────────────────────
def _mht(data: bytes) -> ParseResult:
    import email

    msg = email.message_from_bytes(data)
    parts: list[str] = []
    assets: list[ParsedAsset] = []
    for part in msg.walk():
        ct = part.get_content_type()
        payload = part.get_payload(decode=True) or b""
        if ct.startswith("image/") and payload:
            width, height = _image_size(payload)
            assets.append(ParsedAsset(
                asset_type="image",
                image_bytes=payload,
                media_type=ct,
                width=width,
                height=height,
                needs_ocr=True,
                extraction_model="mhtml-mime",
                extraction_version="v2",
            ))
            continue
        if ct not in ("text/html", "text/plain"):
            continue
        charset = part.get_content_charset() or "utf-8"
        text = payload.decode(charset, "replace")
        parts.append(_strip_html(text) if ct == "text/html" else text)
    return ParseResult(
        "\n".join(p for p in parts if p.strip()).strip(),
        mime_type="multipart/related",
        assets=assets,
    )


# ── HWP ────────────────────────────────────────────────────────
def _hwp(name: str, data: bytes) -> ParseResult:
    """HWP/HWPX 텍스트 추출. 실패 시 빈 텍스트 → 파이프라인에서 failed 처리(§7.H)."""
    try:
        from . import hwp_support
        if name.endswith(".hwpx"):
            return ParseResult(
                hwp_support.extract_hwpx(data),
                mime_type="application/hwpx",
                assets=_zip_media_assets(data, prefix="BinData/"),
            )
        return ParseResult(hwp_support.extract_hwp(data), mime_type="application/hwp")
    except Exception:
        return ParseResult("", needs_ocr=False, mime_type="application/hwp")


# ── Structured/visual helpers ──────────────────────────────────
def _table_to_markdown(rows, *, title: str | None = None) -> str:
    normalized = [[_markdown_cell(cell) for cell in row] for row in rows]
    normalized = [row for row in normalized if any(cell for cell in row)]
    if not normalized:
        return f"# {title}" if title else ""
    width = max(len(row) for row in normalized)
    padded = [row + [""] * (width - len(row)) for row in normalized]
    header = padded[0]
    body = padded[1:]
    lines = []
    if title:
        lines.extend((f"# {title}", ""))
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _markdown_cell(value) -> str:
    text = "" if value is None else str(value)
    return re.sub(r"\s+", " ", text.replace("|", "\\|")).strip()


def _chart_caption(sheet_name: str, chart, chart_no: int) -> str:
    title = ""
    try:
        paragraphs = chart.title.tx.rich.p
        title = " ".join(
            run.t for paragraph in paragraphs for run in paragraph.r if getattr(run, "t", None)
        ).strip()
    except Exception:
        title = ""
    return f"시트: {sheet_name}\n차트 {chart_no}: {title or type(chart).__name__}"


def _zip_media_assets(data: bytes, *, prefix: str) -> list[ParsedAsset]:
    assets: list[ParsedAsset] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for name in sorted(archive.namelist()):
                normalized = name.replace("\\", "/")
                if not normalized.lower().startswith(prefix.lower()) or normalized.endswith("/"):
                    continue
                media_type = _media_type(name)
                if not media_type.startswith("image/"):
                    continue
                payload = archive.read(name)
                width, height = _image_size(payload)
                assets.append(ParsedAsset(
                    asset_type="image",
                    caption=name.rsplit("/", 1)[-1],
                    image_bytes=payload,
                    media_type=media_type,
                    width=width,
                    height=height,
                    needs_ocr=True,
                    extraction_model="office-media",
                    extraction_version="v2",
                ))
    except Exception:
        return []
    return assets


def _media_type(filename: str) -> str:
    guessed = mimetypes.guess_type(filename)[0]
    if guessed in {"image/png", "image/jpeg", "image/webp", "image/gif", "image/tiff", "image/bmp"}:
        return guessed
    return "image/png" if filename.lower().endswith(".png") else "image/jpeg"


def _image_size(data: bytes) -> tuple[int | None, int | None]:
    try:
        from PIL import Image

        with Image.open(io.BytesIO(data)) as image:
            return image.width, image.height
    except Exception:
        return None, None
