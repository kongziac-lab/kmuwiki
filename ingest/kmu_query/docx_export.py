"""DOCX export helpers for Hermes drafts."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def build_approval_docx(
    *,
    title: str,
    body: str,
    source_label: str = "",
    approval_form_plan: list[str] | None = None,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    _configure_styles(doc)
    clean_title = _strip_docx_extension(safe_docx_filename(title))

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("전자결재 문서 초안")
    run.bold = True
    run.font.size = Pt(16)
    _set_run_font(run, "Malgun Gothic")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(clean_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    _set_run_font(title_run, "Malgun Gothic")

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    _set_table_widths(meta, [3.2, 13.4])
    rows = [
        ("문서제목", clean_title),
        ("원문근거", source_label or "검토 필요"),
        ("생성상태", "draft / 사람 검토 필요"),
        ("출력형식", "DOCX 전자결재 결재문 형식"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_cell(cell, "F1F5F9" if cell is row.cells[0] else "FFFFFF")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, "Malgun Gothic")
                    run.font.size = Pt(9.5)

    _add_section_heading(doc, "본문")
    for line in _body_lines(body):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(line)
        run.font.size = Pt(10.5)
        _set_run_font(run, "Malgun Gothic")

    if approval_form_plan:
        _add_section_heading(doc, "생성 기준")
        for item in approval_form_plan:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(item)
            run.font.size = Pt(9.5)
            _set_run_font(run, "Malgun Gothic")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("KMU Wiki draft")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)
    _set_run_font(footer_run, "Malgun Gothic")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def safe_docx_filename(title: str) -> str:
    filename = title.strip().split("/")[-1].split("\\")[-1]
    filename = re.sub(r"[\\/:*?\"<>|]+", "", filename).strip()
    filename = re.sub(r"\.[^.]+$", "", filename)
    filename = filename.strip(". ") or "draft"
    return f"{filename}.docx"


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)


def _add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    _set_run_font(run, "Malgun Gothic")


def _body_lines(body: str) -> list[str]:
    lines = [line.strip() for line in body.splitlines()]
    return [line for line in lines if line] or ["본문 초안 없음"]


def _strip_docx_extension(filename: str) -> str:
    return re.sub(r"\.docx$", "", filename, flags=re.IGNORECASE)


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_table_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
